import io

from docx import Document as DocxDocument

from app.core.exceptions import ValidationException
from app.processing.parsers.base_parser import BaseParser


class DocxParser(BaseParser):
    """Extracts text content from Word (.docx) files."""

    def extract_text(self, file_bytes: bytes) -> str:
        try:
            document = DocxDocument(io.BytesIO(file_bytes))
        except Exception as exc:
            raise ValidationException("Could not read DOCX file — it may be corrupted") from exc

        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

        table_texts = []
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_texts.append(row_text)

        all_text = paragraphs + table_texts
        return "\n".join(all_text).strip()